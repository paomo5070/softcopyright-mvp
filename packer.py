"""
软著工厂 - Word 文档封装器
生成符合软著申报格式的 Word 文档
"""

import argparse
import glob
import json
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn


def set_font(run, font_name="仿宋_GB2312", font_size=12, bold=False):
    """设置字体（中文必须用中文字体名）"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    # 设置中文字体
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def add_page_header_footer(doc, software_name: str, total_pages: int = 30):
    """设置页眉页脚"""
    section = doc.sections[0]

    # 页眉
    header = section.header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_para.add_run(software_name)
    set_font(run, "宋体", 10)

    # 页脚
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run(f"第 1 页 共 {total_pages} 页")
    set_font(run, "宋体", 9)


def create_source_code_doc(code_dir: str, output_path: str, software_name: str):
    """生成源代码文档（前30页+后30页）"""
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '仿宋_GB2312'
    style.font.size = Pt(10.5)  # 五号字
    style.paragraph_format.line_spacing = 1.5
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"{software_name} - 源代码")
    set_font(run, "黑体", 16, bold=True)

    doc.add_paragraph()

    # 收集所有代码文件
    code_files = []
    for root, dirs, files in os.walk(code_dir):
        for f in sorted(files):
            if f.endswith(('.py', '.js', '.ts', '.html', '.css')):
                full_path = os.path.join(root, f)
                rel_path = os.path.relpath(full_path, code_dir)
                code_files.append((rel_path, full_path))

    if not code_files:
        doc.add_paragraph("（未找到源代码文件）")
        doc.save(output_path)
        return

    # 计算总行数，决定截取范围
    all_lines = []
    for rel_path, full_path in code_files:
        with open(full_path, "r", encoding="utf-8") as f:
            lines = f.readlines()
            all_lines.append((rel_path, lines))

    # 前 30 页（每页约 50 行 = 1500 行）
    LINES_PER_PAGE = 50
    FRONT_PAGES = 30
    BACK_PAGES = 30

    total_lines = sum(len(lines) for _, lines in all_lines)

    if total_lines <= (FRONT_PAGES + BACK_PAGES) * LINES_PER_PAGE:
        # 代码量少，全部展示
        for rel_path, lines in all_lines:
            para = doc.add_paragraph()
            run = para.add_run(f"文件：{rel_path}")
            set_font(run, "黑体", 11, bold=True)

            for line in lines:
                para = doc.add_paragraph()
                run = para.add_run(line.rstrip())
                set_font(run, "Courier New", 9)
    else:
        # 截取前 FRONT_PAGES 页
        line_count = 0
        front_limit = FRONT_PAGES * LINES_PER_PAGE

        doc.add_heading("第一部分（前 30 页）", level=2)
        for rel_path, lines in all_lines:
            if line_count >= front_limit:
                break
            para = doc.add_paragraph()
            run = para.add_run(f"文件：{rel_path}")
            set_font(run, "黑体", 11, bold=True)

            for line in lines:
                if line_count >= front_limit:
                    break
                para = doc.add_paragraph()
                run = para.add_run(line.rstrip())
                set_font(run, "Courier New", 9)
                line_count += 1

        # 截取后 BACK_PAGES 页
        doc.add_page_break()
        doc.add_heading("第二部分（后 30 页）", level=2)
        back_limit = BACK_PAGES * LINES_PER_PAGE
        back_start = max(0, total_lines - back_limit)

        line_count = 0
        for rel_path, lines in all_lines:
            for line in lines:
                line_count += 1
                if line_count < back_start:
                    continue
                if line_count >= back_start + back_limit:
                    break
                para = doc.add_paragraph()
                run = para.add_run(line.rstrip())
                set_font(run, "Courier New", 9)

    add_page_header_footer(doc, software_name, total_pages=min(60, total_lines // LINES_PER_PAGE + 1))
    doc.save(output_path)
    print(f"   ✅ 源代码文档: {output_path}")


def create_description_doc(p1_data: dict, p2_data: dict, screenshot_dir: str, output_path: str, software_name: str):
    """生成软件说明书（含截图和功能描述）"""
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = '仿宋_GB2312'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"{software_name}\n软件设计说明书")
    set_font(run, "黑体", 18, bold=True)

    doc.add_paragraph()

    # 1. 软件概述
    doc.add_heading("一、软件概述", level=1)
    profile = p1_data.get("company_profile", {})
    purpose = profile.get("software_purpose", "本软件旨在为用户提供专业的业务管理解决方案。")
    doc.add_paragraph(purpose)

    # 2. 功能说明
    doc.add_heading("二、功能说明", level=1)
    features = p1_data.get("feature_outline", [])
    for i, feat in enumerate(features, 1):
        doc.add_paragraph(f"{i}. {feat}")

    # 3. 运行截图
    doc.add_heading("三、软件运行截图", level=1)
    screenshots = sorted(glob.glob(os.path.join(screenshot_dir, "*.png")))

    if screenshots:
        for i, ss_path in enumerate(screenshots, 1):
            doc.add_paragraph(f"图 {i}：")
            try:
                doc.add_picture(ss_path, width=Cm(14))
            except Exception as e:
                doc.add_paragraph(f"（截图加载失败: {e}）")
            doc.add_paragraph()
    else:
        doc.add_paragraph("（未找到截图文件）")

    # 4. 技术架构
    doc.add_heading("四、技术架构", level=1)
    arch_notes = p2_data.get("architecture_notes", "本软件采用分层架构设计。")
    doc.add_paragraph(arch_notes)

    files = p2_data.get("file_tree", [])
    if files:
        doc.add_paragraph(f"项目包含 {len(files)} 个源文件：")
        for f in files:
            doc.add_paragraph(f"  - {f.get('path', '?')}: {f.get('description', '')}")

    add_page_header_footer(doc, software_name)
    doc.save(output_path)
    print(f"   ✅ 设计说明书: {output_path}")


def main():
    parser = argparse.ArgumentParser(description="软著工厂 - Word 封装器")
    parser.add_argument("--project-dir", required=True, help="项目目录")
    parser.add_argument("--output", required=True, help="输出目录")
    args = parser.parse_args()

    os.makedirs(args.output, exist_ok=True)

    # 加载数据
    p1_path = os.path.join(args.project_dir, "p1_insight.json")
    p2_path = os.path.join(args.project_dir, "p2_architecture.json")

    p1_data = {}
    p2_data = {}
    if os.path.exists(p1_path):
        with open(p1_path, "r", encoding="utf-8") as f:
            p1_data = json.load(f)
    if os.path.exists(p2_path):
        with open(p2_path, "r", encoding="utf-8") as f:
            p2_data = json.load(f)

    software_name = p1_data.get("company_profile", {}).get("name", "软件系统")
    code_dir = os.path.join(args.project_dir, "code")
    screenshot_dir = os.path.join(args.project_dir, "screenshots")

    # 生成源代码文档
    create_source_code_doc(
        code_dir,
        os.path.join(args.output, "源代码文档.docx"),
        software_name,
    )

    # 生成设计说明书
    create_description_doc(
        p1_data, p2_data, screenshot_dir,
        os.path.join(args.output, "软件设计说明书.docx"),
        software_name,
    )

    print(f"\n📦 封装完成: {args.output}")


if __name__ == "__main__":
    main()
